import os
import tkinter as tk
from tkinter import filedialog, messagebox
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinterdnd2 import TkinterDnD, DND_FILES

def markdown_to_word(markdown_file, word_file):
    # 读取 Markdown 文件
    with open(markdown_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 转换 Markdown 为 HTML
    html_content = markdown(md_content)

    # 使用 BeautifulSoup 解析 HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # 创建 Word 文档
    doc = Document()

    # 遍历 HTML 元素并将其转换为 Word 格式
    for element in soup.children:
        if element.name == 'h1':
            add_heading(doc, element.text, level=1)
        elif element.name == 'h2':
            add_heading(doc, element.text, level=2)
        elif element.name == 'h3':
            add_heading(doc, element.text, level=3)
        elif element.name == 'p':
            add_paragraph(doc, element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                add_bullet_point(doc, li.text)
        elif element.name == 'ol':
            for li in element.find_all('li'):
                add_numbered_point(doc, li.text)
        elif element.name == 'code':
            add_code_block(doc, element.text)

    # 保存为 Word 文档
    doc.save(word_file)
    return True

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style.font.size = Pt(12)

def add_bullet_point(doc, text):
    paragraph = doc.add_paragraph(text, style='List Bullet')

def add_numbered_point(doc, text):
    paragraph = doc.add_paragraph(text, style='List Number')

def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def select_and_convert_file():
    # 打开文件选择对话框
    markdown_file = filedialog.askopenfilename(
        title="选择 Markdown 文件",
        filetypes=[("Markdown 文件", "*.md"), ("所有文件", "*.*")]
    )
    if not markdown_file:
        return

    # 确定保存的 Word 文件路径
    word_file = filedialog.asksaveasfilename(
        title="保存为 Word 文件",
        defaultextension=".docx",
        filetypes=[("Word 文件", "*.docx")]
    )
    if not word_file:
        return

    try:
        # 转换 Markdown 到 Word
        markdown_to_word(markdown_file, word_file)
        messagebox.showinfo("转换成功", f"文件已成功转换为 Word 文档：\n{word_file}")
    except Exception as e:
        messagebox.showerror("转换失败", f"文件转换失败：\n{e}")

def drag_and_drop(event):
    file_path = event.data.strip()
    if not file_path.lower().endswith(".md"):
        messagebox.showerror("错误", "请拖动 Markdown 文件 (.md)")
        return

    word_file = filedialog.asksaveasfilename(
        title="保存为 Word 文件",
        defaultextension=".docx",
        filetypes=[("Word 文件", "*.docx")]
    )
    if not word_file:
        return

    try:
        markdown_to_word(file_path, word_file)
        messagebox.showinfo("转换成功", f"文件已成功转换为 Word 文档：\n{word_file}")
    except Exception as e:
        messagebox.showerror("转换失败", f"文件转换失败：\n{e}")

# 创建主界面
from tkinterdnd2 import TkinterDnD, DND_FILES
app = TkinterDnD.Tk()  # 替换为 TkinterDnD 的 Tk 类
app.title("Markdown 转 Word 工具")
app.geometry("500x300")
app.resizable(False, False)

# 标签
label = tk.Label(app, text="拖动 Markdown 文件到此窗口，或点击按钮选择文件", wraplength=400, font=("Arial", 14))
label.pack(pady=30)

# 按钮
button = tk.Button(app, text="选择文件进行转换", font=("Arial", 12), command=select_and_convert_file)
button.pack(pady=20)

# 配置拖放支持
app.drop_target_register(DND_FILES)
app.dnd_bind('<<Drop>>', drag_and_drop)

# 运行主循环
app.mainloop()

